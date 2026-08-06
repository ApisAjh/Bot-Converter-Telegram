"""
Converter PDF: Word→PDF, Image→PDF, PDF→Image, Merge, Split, Compress.

Catatan penting soal fidelity:
- Word → PDF di sini TIDAK memakai LibreOffice (tidak tersedia di Vercel
  serverless), melainkan mengekstrak paragraf & tabel dari .docx lalu
  merender ulang ke PDF memakai reportlab. Formatting kompleks (gambar,
  styling lanjutan) tidak dipertahankan 1:1. Lihat README bagian "Batasan".
- PDF → Image memakai pypdfium2 (murni Python + binary pdfium bawaan pip,
  TIDAK butuh poppler system package) supaya tetap jalan di lingkungan
  serverless seperti Vercel.
"""
from __future__ import annotations

from pathlib import Path

from PyPDF2 import PdfReader, PdfWriter
from PIL import Image
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

import pypdfium2 as pdfium


def word_to_pdf(input_path: Path, output_path: Path) -> None:
    doc = DocxDocument(str(input_path))
    styles = getSampleStyleSheet()
    story = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            story.append(Paragraph(text, styles["Normal"]))
            story.append(Spacer(1, 0.2 * cm))
    for table in doc.tables:
        data = [[cell.text for cell in row.cells] for row in table.rows]
        if data:
            t = Table(data)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.3 * cm))
    if not story:
        story.append(Paragraph("(Dokumen kosong)", styles["Normal"]))
    SimpleDocTemplate(str(output_path), pagesize=A4).build(story)


def image_to_pdf(input_paths: list[Path], output_path: Path) -> None:
    images = [Image.open(p).convert("RGB") for p in input_paths]
    first, rest = images[0], images[1:]
    first.save(str(output_path), save_all=True, append_images=rest)


def pdf_to_images(input_path: Path, output_dir: Path) -> list[Path]:
    pdf = pdfium.PdfDocument(str(input_path))
    output_paths: list[Path] = []
    for i in range(len(pdf)):
        page = pdf[i]
        bitmap = page.render(scale=2.0)
        pil_image = bitmap.to_pil()
        out_path = output_dir / f"page_{i + 1}.png"
        pil_image.save(out_path)
        output_paths.append(out_path)
    pdf.close()
    return output_paths


def merge_pdfs(input_paths: list[Path], output_path: Path) -> None:
    writer = PdfWriter()
    for p in input_paths:
        reader = PdfReader(str(p))
        for page in reader.pages:
            writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)


def _parse_page_spec(spec: str, total_pages: int) -> list[int]:
    """Parse string seperti '1,3,5-7' menjadi list index halaman (0-based)."""
    pages: set[int] = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            start, end = part.split("-", 1)
            start_i, end_i = int(start), int(end)
            pages.update(range(start_i - 1, end_i))
        else:
            pages.add(int(part) - 1)
    return sorted(p for p in pages if 0 <= p < total_pages)


def split_pdf(input_path: Path, output_dir: Path, page_spec: str) -> list[Path]:
    reader = PdfReader(str(input_path))
    total = len(reader.pages)

    if page_spec.strip().lower() == "all":
        indices = list(range(total))
        output_paths = []
        for i in indices:
            writer = PdfWriter()
            writer.add_page(reader.pages[i])
            out_path = output_dir / f"page_{i + 1}.pdf"
            with open(out_path, "wb") as f:
                writer.write(f)
            output_paths.append(out_path)
        return output_paths

    indices = _parse_page_spec(page_spec, total)
    if not indices:
        raise ValueError("Nomor halaman tidak valid atau di luar jangkauan dokumen.")
    writer = PdfWriter()
    for i in indices:
        writer.add_page(reader.pages[i])
    out_path = output_dir / "split_result.pdf"
    with open(out_path, "wb") as f:
        writer.write(f)
    return [out_path]


def compress_pdf(input_path: Path, output_path: Path) -> None:
    reader = PdfReader(str(input_path))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    for page in writer.pages:
        try:
            page.compress_content_streams()
        except Exception:
            pass  # kalau gagal dikompres, tetap lanjut tanpa kompresi stream
    with open(output_path, "wb") as f:
        writer.write(f)
